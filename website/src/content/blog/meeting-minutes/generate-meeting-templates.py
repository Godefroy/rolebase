#!/usr/bin/env python3
"""Generate the Word templates offered on the meeting minutes post.

Lives beside the page it serves: `en.mdx` and `fr.mdx` in this folder,
published on /en/blog/meeting-minutes and /fr/blog/meeting-minutes.

Files produced, under `public/downloads/` and served on `/downloads/<file>`
    meeting-minutes-template.docx        EN, linked from en.mdx
    meeting-agenda-template.docx         EN, linked from en.mdx
    compte-rendu-de-reunion-modele.docx  FR, linked from fr.mdx
    ordre-du-jour-reunion-modele.docx    FR, linked from fr.mdx

These binaries are build artefacts committed to the repo. Edit the definitions
below and re-run this script rather than editing the files themselves.

    pip install python-docx
    python3 src/content/blog/meeting-minutes/generate-meeting-templates.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parents[4] / "public" / "downloads"

PURPLE = RGBColor(0x98, 0x70, 0xF0)
GREY = RGBColor(0x6B, 0x6B, 0x6B)


def styled(doc):
    """Base typography shared by every template."""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for name, size in (("Heading 1", 20), ("Heading 2", 14)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = PURPLE
        st.font.bold = True
    return doc


def field(doc, lang, label, hint=""):
    """A `Label:` line with a greyed-out hint the writer overwrites."""
    p = doc.add_paragraph()
    run = p.add_run(f"{label} : " if lang == "fr" else f"{label}: ")
    run.bold = True
    if hint:
        h = p.add_run(hint)
        h.italic = True
        h.font.color.rgb = GREY
    return p


def blanks(doc, n, style=None):
    for _ in range(n):
        doc.add_paragraph("", style=style)


def actions_table(doc, headers, rows=4):
    table = doc.add_table(rows=1 + rows, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    return table


def footer_note(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREY


def minutes(t, lang):
    doc = styled(Document())
    doc.add_heading(t["title"], level=1)
    for label, hint in t["fields"]:
        field(doc, lang, label, hint)
    doc.add_heading(t["agenda"], level=2)
    for i in range(1, 4):
        doc.add_paragraph(f"{i}. ")
    doc.add_heading(t["discussion"], level=2)
    for n in (1, 2):
        doc.add_heading(f"{t['item']} {n}", level=3)
        for label in t["blocks"]:
            field(doc, lang, label)
    doc.add_heading(t["actions"], level=2)
    actions_table(doc, t["action_cols"])
    doc.add_paragraph()
    doc.add_heading(t["next"], level=2)
    for label in t["next_fields"]:
        field(doc, lang, label)
    footer_note(doc, t["footer"])
    return doc


def agenda(t, lang):
    doc = styled(Document())
    doc.add_heading(t["title"], level=1)
    for label, hint in t["fields"]:
        field(doc, lang, label, hint)
    doc.add_paragraph()
    actions_table(doc, t["cols"], rows=6)
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(t["tip"])
    run.italic = True
    run.font.color.rgb = GREY
    footer_note(doc, t["footer"])
    return doc


FR_FOOTER = "Modèle proposé par Rolebase, rolebase.io"
EN_FOOTER = "Template by Rolebase, rolebase.io"

TEMPLATES = {
    "fr": {
        "compte-rendu-de-reunion-modele.docx": (minutes, {
            "title": "Compte rendu de réunion",
            "fields": [("Réunion", "nom de la réunion"), ("Date", ""), ("Horaire", ""),
                       ("Lieu ou lien", ""), ("Participants", ""), ("Excusés", ""),
                       ("Compte rendu rédigé par", "")],
            "agenda": "Ordre du jour",
            "discussion": "Discussions et décisions",
            "item": "Point",
            "blocks": ["Discuté", "Décidé", "Pourquoi"],
            "actions": "Actions",
            "action_cols": ["Action", "Responsable", "Échéance"],
            "next": "Prochaine réunion",
            "next_fields": ["Date", "Reporté"],
            "footer": FR_FOOTER,
        }),
        "ordre-du-jour-reunion-modele.docx": (agenda, {
            "title": "Ordre du jour",
            "fields": [("Réunion", "nom de la réunion"), ("Date", ""),
                       ("Objet de cette réunion", "ce que la réunion doit produire"),
                       ("Participants", ""), ("Durée", "")],
            "cols": ["Durée", "Point", "Proposé par", "Résultat attendu"],
            "tip": "Un point sans résultat attendu se transforme en conversation de vingt minutes. "
                   "Remplissez la colonne « Résultat attendu » avant d'envoyer l'ordre du jour.",
            "footer": FR_FOOTER,
        }),
    },
    "en": {
        "meeting-minutes-template.docx": (minutes, {
            "title": "Meeting minutes",
            "fields": [("Meeting", "name of the meeting"), ("Date", ""), ("Time", ""),
                       ("Location or link", ""), ("Attendees", ""), ("Apologies", ""),
                       ("Minutes taken by", "")],
            "agenda": "Agenda",
            "discussion": "Discussion and decisions",
            "item": "Item",
            "blocks": ["Discussed", "Decided", "Why"],
            "actions": "Actions",
            "action_cols": ["Action", "Owner", "Due date"],
            "next": "Next meeting",
            "next_fields": ["Date", "Carried over"],
            "footer": EN_FOOTER,
        }),
        "meeting-agenda-template.docx": (agenda, {
            "title": "Meeting agenda",
            "fields": [("Meeting", "name of the meeting"), ("Date", ""),
                       ("Purpose of this meeting", "what the meeting has to produce"),
                       ("Attendees", ""), ("Duration", "")],
            "cols": ["Time", "Item", "Brought by", "Outcome needed"],
            "tip": "An item with no stated outcome turns into a twenty-minute conversation. "
                   "Fill in the Outcome needed column before sending the agenda out.",
            "footer": EN_FOOTER,
        }),
    },
}

if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    for lang, files in TEMPLATES.items():
        for filename, (builder, spec) in files.items():
            builder(spec, lang).save(OUT / filename)
            print(f"  {lang}  {filename}")
