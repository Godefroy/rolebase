#!/usr/bin/env python3
"""Generate the Word templates offered on the roles and responsibilities post.

Lives beside the page it serves: `en.mdx` and `fr.mdx` in this folder,
published on /en/blog/roles-responsibilities-template and
/fr/blog/roles-responsibilities-template.

Files produced, under `public/downloads/` and served on `/downloads/<file>`
    roles-responsibilities-template.docx  EN, linked from en.mdx
    modele-roles-responsabilites.docx     FR, linked from fr.mdx

Each document carries a short instruction note, three blank role pages built on
the six fields of the article, a team sheet listing every role on one line, and
a filled example page the reader deletes.

These binaries are build artefacts committed to the repo. Edit the definitions
below and re-run this script rather than editing the files themselves.

    pip install python-docx
    python3 src/content/blog/roles-responsibilities-template/generate-roles-template.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parents[4] / "public" / "downloads"

PURPLE = RGBColor(0x98, 0x70, 0xF0)
GREY = RGBColor(0x6B, 0x6B, 0x6B)

BLANK_ROLE_PAGES = 3
ACCOUNTABILITY_LINES = 5
TEAM_SHEET_ROWS = 12


def styled(doc):
    """Base typography shared by every template."""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for name, size in (("Heading 1", 20), ("Heading 2", 14)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = PURPLE
        st.font.bold = True
    return doc


def hint(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = GREY
    return p


def field(doc, lang, label, note="", value=""):
    """A `Label:` line, followed either by a greyed hint or by a filled value."""
    p = doc.add_paragraph()
    run = p.add_run(f"{label} : " if lang == "fr" else f"{label}: ")
    run.bold = True
    if value:
        p.add_run(value)
    elif note:
        h = p.add_run(note)
        h.italic = True
        h.font.color.rgb = GREY
    return p


def role_page(doc, t, lang, filled=None):
    """One role per page: the six fields, with the accountabilities as a list."""
    doc.add_heading(t["role_heading"] if filled is None else t["example_heading"], level=1)
    labels = t["fields"]
    field(doc, lang, labels["name"][0], labels["name"][1],
          filled["name"] if filled else "")
    field(doc, lang, labels["purpose"][0], labels["purpose"][1],
          filled["purpose"] if filled else "")
    field(doc, lang, labels["domain"][0], labels["domain"][1],
          filled["domain"] if filled else "")

    p = doc.add_paragraph()
    run = p.add_run(f"{labels['accountabilities'][0]} : " if lang == "fr"
                    else f"{labels['accountabilities'][0]}: ")
    run.bold = True
    h = p.add_run(labels["accountabilities"][1])
    h.italic = True
    h.font.color.rgb = GREY
    lines = filled["accountabilities"] if filled else [""] * ACCOUNTABILITY_LINES
    for line in lines:
        doc.add_paragraph(line, style="List Bullet")

    field(doc, lang, labels["holder"][0], labels["holder"][1],
          filled["holder"] if filled else "")
    field(doc, lang, labels["reviewed"][0], labels["reviewed"][1],
          filled["reviewed"] if filled else "")


def team_sheet(doc, t):
    doc.add_heading(t["team_heading"], level=1)
    hint(doc, t["team_hint"])
    table = doc.add_table(rows=1 + TEAM_SHEET_ROWS, cols=len(t["team_cols"]))
    table.style = "Table Grid"
    for i, header in enumerate(t["team_cols"]):
        cell = table.rows[0].cells[i]
        cell.text = ""
        cell.paragraphs[0].add_run(header).bold = True


def footer_note(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREY


def build(t, lang):
    doc = styled(Document())
    doc.add_heading(t["title"], level=1)
    hint(doc, t["intro"])
    for _ in range(BLANK_ROLE_PAGES):
        doc.add_page_break()
        role_page(doc, t, lang)
    doc.add_page_break()
    team_sheet(doc, t)
    doc.add_page_break()
    role_page(doc, t, lang, filled=t["example"])
    hint(doc, t["example_hint"])
    footer_note(doc, t["footer"])
    return doc


FR_FOOTER = "Modèle proposé par Rolebase, rolebase.io"
EN_FOOTER = "Template by Rolebase, rolebase.io"

TEMPLATES = {
    "en": ("roles-responsibilities-template.docx", {
        "title": "Roles and responsibilities",
        "intro": "One page per role, filled in with the person who holds it. Copy a blank "
                 "page for every new role, and list them all on the team sheet.",
        "role_heading": "Role",
        "example_heading": "Filled example",
        "team_heading": "Team sheet",
        "team_hint": "One line per role, so the whole team is readable on one page.",
        "team_cols": ["Role", "Held by", "Purpose in one line", "Last reviewed"],
        "fields": {
            "name": ("Role name", "a responsibility, two or three words"),
            "purpose": ("Purpose", "one sentence, no deadline, no number"),
            "domain": ("Decides alone on", "the assets and processes this role controls"),
            "accountabilities": ("Accountabilities", "three to seven, each starting with an -ing verb"),
            "holder": ("Held by", "a name, or vacant"),
            "reviewed": ("Last reviewed", "a date"),
        },
        "example": {
            "name": "Editorial Content",
            "purpose": "Readers who find their answer on our site before they talk to anyone",
            "domain": "The publication calendar and the style guide",
            "accountabilities": [
                "Publishing to the quarterly calendar",
                "Briefing and reviewing external writers",
                "Reporting readership at the monthly team meeting",
                "Retiring pages that no longer match the product",
            ],
            "holder": "Amina",
            "reviewed": "12 June",
        },
        "example_hint": "Delete this page once your own roles are written.",
        "footer": EN_FOOTER,
    }),
    "fr": ("modele-roles-responsabilites.docx", {
        "title": "Rôles et responsabilités",
        "intro": "Une page par rôle, remplie avec la personne qui le tient. Copiez une page "
                 "vierge pour chaque nouveau rôle, et reportez-les tous sur la feuille d'équipe.",
        "role_heading": "Rôle",
        "example_heading": "Exemple rempli",
        "team_heading": "Feuille d'équipe",
        "team_hint": "Une ligne par rôle, pour lire toute l'équipe sur une page.",
        "team_cols": ["Rôle", "Tenu par", "Raison d'être en une ligne", "Dernière relecture"],
        "fields": {
            "name": ("Nom du rôle", "une responsabilité, deux ou trois mots"),
            "purpose": ("Raison d'être", "une phrase, sans échéance ni chiffre"),
            "domain": ("Décide seul de", "les actifs et les processus que ce rôle contrôle"),
            "accountabilities": ("Redevabilités", "trois à sept, chacune commençant par un infinitif"),
            "holder": ("Tenu par", "un prénom, ou vacant"),
            "reviewed": ("Dernière relecture", "une date"),
        },
        "example": {
            "name": "Contenu éditorial",
            "purpose": "Des lecteurs qui trouvent leur réponse sur notre site avant de parler à quelqu'un",
            "domain": "Le calendrier de publication et la charte éditoriale",
            "accountabilities": [
                "Publier selon le calendrier trimestriel",
                "Cadrer et relire les rédacteurs externes",
                "Présenter l'audience à la réunion mensuelle",
                "Retirer les pages qui ne correspondent plus au produit",
            ],
            "holder": "Amina",
            "reviewed": "12 juin",
        },
        "example_hint": "Supprimez cette page une fois vos propres rôles écrits.",
        "footer": FR_FOOTER,
    }),
}

if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    for lang, (filename, spec) in TEMPLATES.items():
        build(spec, lang).save(OUT / filename)
        print(f"  {lang}  {filename}")
